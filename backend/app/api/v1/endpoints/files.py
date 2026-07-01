import hashlib
import shutil
from typing import IO
from fastapi import APIRouter, File, UploadFile, HTTPException, Depends, Form
from fastapi.responses import FileResponse
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
from stegano import lsb
import os
from app.api import deps
from app.models.user import User
from sqlmodel import Session
from datetime import datetime, timezone

def parse_pdf_date(date_str: str) -> datetime | None:
    """Parse PDF date format like D:20220101123456Z or D:20220101123456+05'30'"""
    if not date_str:
        return None
    try:
        clean_date = date_str.replace("D:", "").replace("'", "")
        # Handle some variations
        if "+" in clean_date:
            dt_part, tz_part = clean_date.split("+")
            dt = datetime.strptime(dt_part, "%Y%m%d%H%M%S")
            # For simplicity, ignoring precise timezone offset parsing in this quick impl
            # better to stick to UTC or basic parsing if simple
            return dt
        elif "Z" in clean_date:
             return datetime.strptime(clean_date.replace("Z",""), "%Y%m%d%H%M%S")
        else:
             return datetime.strptime(clean_date, "%Y%m%d%H%M%S")
    except Exception:
        return None

router = APIRouter()

TEMP_DIR = "temp_files"
os.makedirs(TEMP_DIR, exist_ok=True)

def calculate_hash(file_path: str, algorithm: str = "sha256") -> str:
    hash_func = getattr(hashlib, algorithm)()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_func.update(chunk)
    return hash_func.hexdigest()

def calculate_md5(file_path: str) -> str:
    hash_func = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_func.update(chunk)
    return hash_func.hexdigest()

@router.post("/integrity")
async def check_file_integrity(
    file: UploadFile = File(...), 
    algorithm: str = "sha256",
    current_user: User = Depends(deps.get_current_user),
    session: Session = Depends(deps.get_session) # Add session dependency
):
    import json
    from app.models.scan_history import ScanHistory
    from sqlmodel import select
    from sqlalchemy import desc
    import datetime

    file_location = f"{TEMP_DIR}/{file.filename}"
    with open(file_location, "wb+") as file_object:
        shutil.copyfileobj(file.file, file_object)
    
    file_hash = calculate_hash(file_location, algorithm)
    md5_hash = calculate_md5(file_location)
    file_stats = os.stat(file_location)
    size_bytes = file_stats.st_size
    os.remove(file_location)
    
    # 1. Fetch Last Baseline (Previous Scan of this file)
    # We look for the most recent scan of SAME type and SAME target by this user
    query = select(ScanHistory).where(
        ScanHistory.user_id == current_user.id,
        ScanHistory.scan_type == "FILE_INTEGRITY",
        ScanHistory.target == file.filename
    ).order_by(desc(ScanHistory.created_at))
    
    last_scan = session.exec(query).first()
    
    baseline_hash = None
    baseline_time = None
    status = "NEW"
    details = "First scan for this file. Baseline created."
    
    if last_scan:
        try:
            last_result = json.loads(last_scan.result)
            # The previous result might be valid, let's extract hash
            # If our previous format was simple, handle that.
            if "hash" in last_result:
                baseline_hash = last_result["hash"]
                baseline_time = last_scan.created_at
                
                if baseline_hash == file_hash:
                    status = "CLEAN"
                    details = "File matches the previous baseline."
                else:
                    status = "MODIFIED"
                    details = "File content differs from previous scan."
        except Exception:
            # Maybe legacy format or parse error, treat as new
            pass

    # Create naive IST timestamp (UTC+5:30) to bypass browser timezone conversion issues
    import datetime
    from datetime import timedelta, timezone
    
    def to_ist_str(dt_source) -> str | None:
        if not dt_source:
            return None
        # Ensure we have a datetime object
        if isinstance(dt_source, (int, float)):
             dt_source = datetime.datetime.fromtimestamp(dt_source, timezone.utc)
        
        # If naive, assume UTC (standard backend practice)
        if dt_source.tzinfo is None:
            dt_source = dt_source.replace(tzinfo=timezone.utc)
            
        ist_time = dt_source.astimezone(timezone(timedelta(hours=5, minutes=30)))
        return ist_time.strftime("%Y-%m-%dT%H:%M:%S")

    scan_time_str = to_ist_str(datetime.datetime.now(timezone.utc))
    
    # Process baseline_time
    baseline_time_str = to_ist_str(baseline_time) if baseline_time else None

    result = {
        "filename": file.filename,
        "algorithm": algorithm, 
        "hash": file_hash, # Current Hash
        "size_bytes": size_bytes,
        "scan_time": scan_time_str,
        # FIM Details
        "status": status,
        "details": details,
        "baseline_hash": baseline_hash,
        "baseline_time": baseline_time_str,
        "md5": md5_hash
    }

    # Save to history
    scan = ScanHistory(
        user_id=current_user.id,
        scan_type="FILE_INTEGRITY",
        target=file.filename,
        result=json.dumps(result, default=str)
    )
    session.add(scan)
    session.commit()

    return result

@router.post("/compare")
async def compare_files(
    file1: UploadFile = File(...),
    file2: UploadFile = File(...),
    current_user: User = Depends(deps.get_current_user),
    session: Session = Depends(deps.get_session)
):
    """Compare two files and return detailed comparison results"""
    import json
    from app.models.scan_history import ScanHistory
    import datetime
    from datetime import timedelta, timezone
    
    file1_location = f"{TEMP_DIR}/compare1_{file1.filename}"
    file2_location = f"{TEMP_DIR}/compare2_{file2.filename}"
    
    try:
        # Save both files
        with open(file1_location, "wb+") as f:
            shutil.copyfileobj(file1.file, f)
        with open(file2_location, "wb+") as f:
            shutil.copyfileobj(file2.file, f)
        
        # Calculate hashes for both files
        file1_sha256 = calculate_hash(file1_location, "sha256")
        file2_sha256 = calculate_hash(file2_location, "sha256")
        file1_md5 = calculate_md5(file1_location)
        file2_md5 = calculate_md5(file2_location)
        
        # Get file sizes
        file1_size = os.stat(file1_location).st_size
        file2_size = os.stat(file2_location).st_size
        
        # Determine match status
        hashes_match = file1_sha256 == file2_sha256
        sizes_match = file1_size == file2_size
        
        if hashes_match:
            status = "IDENTICAL"
            details = "Both files are identical (same content and hash)."
        elif sizes_match:
            status = "MODIFIED"
            details = "Files have the same size but different content."
        else:
            status = "DIFFERENT"
            details = "Files have different sizes and content."
        
        # Calculate size difference
        size_diff = abs(file1_size - file2_size)
        size_diff_pct = (size_diff / max(file1_size, file2_size) * 100) if max(file1_size, file2_size) > 0 else 0
        
        def to_ist_str(dt_source) -> str | None:
            if not dt_source:
                return None
            if isinstance(dt_source, (int, float)):
                dt_source = datetime.datetime.fromtimestamp(dt_source, timezone.utc)
            if dt_source.tzinfo is None:
                dt_source = dt_source.replace(tzinfo=timezone.utc)
            ist_time = dt_source.astimezone(timezone(timedelta(hours=5, minutes=30)))
            return ist_time.strftime("%Y-%m-%dT%H:%M:%S")
        
        scan_time_str = to_ist_str(datetime.datetime.now(timezone.utc))
        
        result = {
            "status": status,
            "details": details,
            "hashes_match": hashes_match,
            "sizes_match": sizes_match,
            "scan_time": scan_time_str,
            "file1": {
                "filename": file1.filename,
                "size_bytes": file1_size,
                "sha256": file1_sha256,
                "md5": file1_md5
            },
            "file2": {
                "filename": file2.filename,
                "size_bytes": file2_size,
                "sha256": file2_sha256,
                "md5": file2_md5
            },
            "comparison": {
                "size_difference_bytes": size_diff,
                "size_difference_percent": round(size_diff_pct, 2)
            }
        }
        
        # Save to history
        scan = ScanHistory(
            user_id=current_user.id,
            scan_type="FILE_COMPARE",
            target=f"{file1.filename} vs {file2.filename}",
            result=json.dumps(result, default=str)
        )
        session.add(scan)
        session.commit()
        
        return result
        
    finally:
        # Cleanup temp files
        if os.path.exists(file1_location):
            os.remove(file1_location)
        if os.path.exists(file2_location):
            os.remove(file2_location)

@router.post("/metadata")
async def analyze_metadata(
    file: UploadFile = File(...),
    last_modified: int = Form(default=None), # Client-provided timestamp (ms)
    current_user: User = Depends(deps.get_current_user),
    session: Session = Depends(deps.get_session)
):
    print(f"DEBUG: Received last_modified={last_modified}")
    # Create temp file
    file_location = f"{TEMP_DIR}/{file.filename}"
    with open(file_location, "wb+") as file_object:
        shutil.copyfileobj(file.file, file_object)
        
    try:
        # 1. Basic File System Metadata
        file_stats = os.stat(file_location)
        
        # Default dates to "now" (creation of temp file)
        created_at = file_stats.st_ctime
        modified_at = file_stats.st_mtime
        
        # Override with client provided modified time if available
        if last_modified:
            # Client sends milliseconds, convert to seconds
            modified_at = last_modified / 1000.0
            created_at = modified_at

        # Helper to convert to naive IST string
        def to_ist_str(dt_source) -> str | None:
            if not dt_source:
                return None
            import datetime
            from datetime import timedelta, timezone
            if isinstance(dt_source, (int, float)):
                 dt_source = datetime.datetime.fromtimestamp(dt_source, timezone.utc)
            if isinstance(dt_source, str): # Already string, maybe ISO? return as is or parsing?
                return dt_source
            if dt_source.tzinfo is None:
                dt_source = dt_source.replace(tzinfo=timezone.utc)
            ist_time = dt_source.astimezone(timezone(timedelta(hours=5, minutes=30)))
            return ist_time.strftime("%Y-%m-%dT%H:%M:%S")

        metadata = {
            "filename": file.filename,
            "content_type": file.content_type,
            "size_bytes": file_stats.st_size,
            "created_at": created_at, # Keep internal types for now, will override at end or use formatting?
            # actually frontend expects string for display or number for Date? 
            # Frontend uses new Date(timestamp * 1000). 
            # If I return a String, new Date(string) works.
            # But "created_at": created_at (which is float)
            # PROPOSAL: Add formatted fields or overwrite?
            # Overwriting changes type. Frontend TS might complain if strictly typed.
            # MetadataViewer uses `any` for result.
            "created_at": to_ist_str(created_at),
            "modified_at": to_ist_str(modified_at),
        }
        
        ext = os.path.splitext(file.filename)[1].lower()
        
        # Helper to update creation time if a better source is found
        def update_creation_time(dt_obj):
            if dt_obj:
                try:
                    # Convert to string first if it's already a string, but we want consistent IST formatting
                    # using to_ist_str handles floats, ints, datetime. 
                    # If it's a raw string (e.g. from docx), we might need to parse it or just trust it?
                    # The previous logic handled "some libs return strings".
                    # Let's try to parse if possible or pass to to_ist_str
                    
                    if isinstance(dt_obj, str):
                        # Attempt to parse if it looks like ISO
                        # But to_ist_str expects datetime or float/int.
                        # If we pass string to to_ist_str as implemented above, it returns valid string.
                        # Wait, the implementation of to_ist_str above: if str return str.
                        # This assumes the string is ALREADY acceptable.
                        # But DOCX 'created' might be specific format.
                        # Let's update metadata if we can.
                        metadata["created_at"] = str(dt_obj)

                    if isinstance(dt_obj, datetime):
                        metadata["created_at"] = to_ist_str(dt_obj)
                        
                except Exception as e:
                    print(f"DEBUG: Failed to update creation time: {e}")

        # 2. Image Metadata
        if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp']:
            print("DEBUG: Processing Image")
            try:
                image = Image.open(file_location)
                metadata["is_image"] = True
                metadata["format"] = image.format
                metadata["mode"] = image.mode
                metadata["dimensions"] = image.size
                
                # Extract Image Info
                img_info = {}
                for k, v in image.info.items():
                    if k != "exif" and k != "photoshop" and isinstance(k, str):
                        img_info[k] = str(v)
                if img_info:
                    metadata["image_info"] = img_info

                # Extract EXIF
                exif_info = {}
                gps_info = {}
                
                exif = image.getexif()
                if exif:
                    # Standard Tags
                    for tag_id, value in exif.items():
                        tag = TAGS.get(tag_id, tag_id)
                        exif_info[str(tag)] = str(value)
                        
                        # Try to find original creation date in EXIF
                        if tag == "DateTimeOriginal" or tag == "DateTime":
                            try:
                                dt_obj = datetime.strptime(str(value), "%Y:%m:%d %H:%M:%S")
                                update_creation_time(dt_obj)
                            except Exception:
                                pass
                    
                    # GPS Tags (IFD 0x8825)
                    gps_ifd = exif.get_ifd(0x8825)
                    if gps_ifd:
                        for tag_id, value in gps_ifd.items():
                            tag = GPSTAGS.get(tag_id, tag_id)
                            gps_info[str(tag)] = str(value)

                if exif_info:
                    metadata["exif"] = exif_info
                if gps_info:
                    metadata["gps"] = gps_info
            except Exception as e:
                print(f"DEBUG: Image processing error: {e}")
                metadata["is_image"] = False

        # 3. PDF Metadata
        elif ext == '.pdf':
            print("DEBUG: Processing PDF")
            try:
                import PyPDF2
                print("DEBUG: Imported PyPDF2")
                with open(file_location, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    pdf_info = pdf.metadata
                    metadata["extended_metadata"] = {
                        "Pages": len(pdf.pages),
                        "Encrypted": pdf.is_encrypted,
                    }
                    if pdf_info:
                        for k, v in pdf_info.items():
                            key = k.replace('/', '')
                            metadata["extended_metadata"][key] = str(v)
                            if key == "CreationDate":
                                dt = parse_pdf_date(str(v))
                                update_creation_time(dt)
                                if dt:
                                    metadata["extended_metadata"][key] = dt.strftime("%Y-%m-%d %H:%M:%S")
                            
                            if key == "ModDate":
                                dt = parse_pdf_date(str(v))
                                if dt:
                                     metadata["extended_metadata"][key] = dt.strftime("%Y-%m-%d %H:%M:%S")
                                     metadata["modified_at"] = dt.timestamp()
                    
                        # Add more detailed PDF info by checking keys directly or relying on loop above
                        # Since we iterate all items above, we don't strictly need these specific checks if we want ALL metadata
                        # But for explicit well-known fields mapping to nice names:
                        if "/Producer" in pdf_info:
                            metadata["extended_metadata"]["Producer"] = str(pdf_info["/Producer"])
                        if "/Creator" in pdf_info:
                            metadata["extended_metadata"]["Creator"] = str(pdf_info["/Creator"])
                        if "/Subject" in pdf_info:
                            metadata["extended_metadata"]["Subject"] = str(pdf_info["/Subject"])
                        if "/Keywords" in pdf_info:
                            metadata["extended_metadata"]["Keywords"] = str(pdf_info["/Keywords"])

                    print(f"DEBUG: PDF Metadata: {metadata['extended_metadata']}")
            except Exception as e:
                print(f"DEBUG: PDF processing error: {e}")
                import traceback
                traceback.print_exc()

        # 4. DOCX Metadata
        elif ext == '.docx':
            print("DEBUG: Processing DOCX")
            try:
                import docx
                from docx.shared import Inches
                doc = docx.Document(file_location)
                core_props = doc.core_properties
                
                # Count document statistics
                paragraph_count = len(doc.paragraphs)
                table_count = len(doc.tables)
                section_count = len(doc.sections)
                
                # Count words and characters
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                all_text = '\n'.join(full_text)
                word_count = len(all_text.split())
                char_count = len(all_text)
                char_count_no_spaces = len(all_text.replace(' ', '').replace('\n', ''))
                
                # Count images (inline shapes)
                image_count = 0
                for para in doc.paragraphs:
                    for run in para.runs:
                        if run._element.xml.find('drawing') != -1:
                            image_count += 1
                
                # Count hyperlinks
                hyperlink_count = 0
                for para in doc.paragraphs:
                    for hyperlink in para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink'):
                        hyperlink_count += 1
                
                # Get styles info
                styles_in_use = set()
                for para in doc.paragraphs:
                    if para.style:
                        styles_in_use.add(para.style.name)
                
                metadata["extended_metadata"] = {
                    # Core Properties
                    "Title": core_props.title,
                    "Author": core_props.author,
                    "Subject": core_props.subject,
                    "Keywords": core_props.keywords,
                    "Category": core_props.category,
                    "Comments": core_props.comments,
                    "Created": str(core_props.created) if core_props.created else None,
                    "Modified": str(core_props.modified) if core_props.modified else None,
                    "Last Modified By": core_props.last_modified_by,
                    "Revision": core_props.revision,
                    "Content Status": core_props.content_status,
                    "Language": core_props.language,
                    "Version": core_props.version,
                }
                
                # Document Statistics (separate section)
                metadata["document_stats"] = {
                    "Paragraphs": paragraph_count,
                    "Word Count": word_count,
                    "Character Count (with spaces)": char_count,
                    "Character Count (no spaces)": char_count_no_spaces,
                    "Tables": table_count,
                    "Sections": section_count,
                    "Images": image_count,
                    "Hyperlinks": hyperlink_count,
                    "Styles Used": len(styles_in_use),
                    "Style Names": ", ".join(sorted(styles_in_use)[:10]) if styles_in_use else None,
                }
                
                update_creation_time(core_props.created)
            except Exception as e:
                 print(f"DEBUG: DOCX processing error: {e}")
                 import traceback
                 traceback.print_exc()
        
        # 5. XLSX Metadata
        elif ext == '.xlsx':
             print("DEBUG: Processing XLSX")
             try:
                import openpyxl
                # Use data_only=False to detect formulas
                wb = openpyxl.load_workbook(file_location, read_only=False, data_only=False)
                props = wb.properties
                
                # Core properties
                metadata["extended_metadata"] = {
                    "Title": props.title,
                    "Author": props.creator,
                    "Subject": props.subject,
                    "Keywords": props.keywords,
                    "Category": props.category,
                    "Description": props.description,
                    "Created": str(props.created) if props.created else None,
                    "Modified": str(props.modified) if props.modified else None,
                    "Last Modified By": props.lastModifiedBy,
                    "Company": props.company,
                    "Manager": props.manager,
                }
                
                # Workbook Statistics
                sheet_info = []
                total_rows = 0
                total_cells = 0
                total_formulas = 0
                
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    row_count = ws.max_row if ws.max_row else 0
                    col_count = ws.max_column if ws.max_column else 0
                    cell_count = row_count * col_count
                    total_rows += row_count
                    total_cells += cell_count
                    
                    # Count formulas in this sheet (sample first 500 cells)
                    formula_count = 0
                    for row in ws.iter_rows(max_row=min(100, row_count), max_col=min(20, col_count)):
                        for cell in row:
                            if cell.value and str(cell.value).startswith('='):
                                formula_count += 1
                    total_formulas += formula_count
                    
                    sheet_info.append(f"{sheet_name} ({row_count}x{col_count})")
                
                metadata["workbook_stats"] = {
                    "Sheet Count": len(wb.sheetnames),
                    "Sheets": ", ".join(sheet_info),
                    "Total Rows (approx)": total_rows,
                    "Total Cells (approx)": total_cells,
                    "Formulas Detected": total_formulas if total_formulas > 0 else "None detected",
                    "Has Named Ranges": len(wb.defined_names.definedName) > 0 if wb.defined_names else False,
                }
                
                update_creation_time(props.created)
                wb.close()
             except Exception as e:
                print(f"DEBUG: XLSX processing error: {e}")
                import traceback
                traceback.print_exc()

        # 6. Audio/Video Metadata
        elif ext in ['.mp3', '.mp4', '.m4a', '.flac', '.ogg', '.wav', '.mkv', '.avi']:
            print("DEBUG: Processing Media")
            try:
                import mutagen
                f = mutagen.File(file_location)
                if f:
                    audio_meta = {
                        "MIME Type": f.mime[0] if f.mime else "Unknown",
                    }
                    if f.info:
                        audio_meta["Duration (s)"] = f"{f.info.length:.2f}"
                        if hasattr(f.info, 'bitrate'):
                            audio_meta["Bitrate"] = f"{f.info.bitrate // 1000} kbps"
                        if hasattr(f.info, 'sample_rate'):
                            audio_meta["Sample Rate"] = f"{f.info.sample_rate} Hz"
                        if hasattr(f.info, 'channels'):
                            audio_meta["Channels"] = f.info.channels
                    
                    # Extract tags (Title, Artist, Album)
                    # Mutagen structure varies by format (ID3, MP4, etc.)
                    # We try to dump generic tags
                    for k, v in f.tags.items() if f.tags else []:
                        if isinstance(v, list):
                            v = ", ".join([str(item) for item in v])
                        # Filter out binary data or very long fields (like lyrics/pictures)
                        if isinstance(v, str) and len(v) < 200:
                            audio_meta[k] = v
                            
                    metadata["extended_metadata"] = audio_meta
            except Exception as e:
                print(f"DEBUG: Media processing error: {e}")

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error analyzing file: {str(e)}")
    finally:
        if os.path.exists(file_location):
            os.remove(file_location)

    # Save to history
    import json
    from app.models.scan_history import ScanHistory
    
    scan = ScanHistory(
        user_id=current_user.id,
        scan_type="METADATA_SCAN",
        target=file.filename,
        result=json.dumps(metadata, default=str)
    )
    session.add(scan)
    session.commit()
            
    return metadata

@router.post("/stegano/encode")
async def stegano_encode(
    message: str, 
    file: UploadFile = File(...),
    current_user: User = Depends(deps.get_current_user)
):
    file_location = f"{TEMP_DIR}/{file.filename}"
    output_location = f"{TEMP_DIR}/stegano_{file.filename}"
    with open(file_location, "wb+") as file_object:
        shutil.copyfileobj(file.file, file_object)
        
    try:
        secret = lsb.hide(file_location, message)
        secret.save(output_location)
    except Exception as e:
        if os.path.exists(file_location):
            os.remove(file_location)
        raise HTTPException(status_code=400, detail=f"Steganography failed: {str(e)}")

    os.remove(file_location)
    return FileResponse(output_location, filename=f"stegano_{file.filename}") # Cleanup needed in production

@router.post("/stegano/decode")
async def stegano_decode(
    file: UploadFile = File(...),
    current_user: User = Depends(deps.get_current_user)
):
    file_location = f"{TEMP_DIR}/{file.filename}"
    with open(file_location, "wb+") as file_object:
        shutil.copyfileobj(file.file, file_object)
        
    try:
        clear_message = lsb.reveal(file_location)
    except Exception as e:
        if os.path.exists(file_location):
            os.remove(file_location)
        raise HTTPException(status_code=400, detail=f"Could not decode message: {str(e)}")
        
    os.remove(file_location)
    return {"message": clear_message}
